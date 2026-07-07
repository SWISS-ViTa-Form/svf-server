from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import io, os, subprocess, tempfile, base64, requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)
CORS(app, origins=['https://portail-swissvitaform.netlify.app', 'http://localhost:3000', '*'])

CERT_COMPLET = base64.b64decode(open('/app/cert_complet.b64').read())
CERT_COMPACT = base64.b64decode(open('/app/cert_compact.b64').read())

def clear_para(para):
    for run in para.runs:
        run.text = ''

def set_center(para):
    pPr = para._p.get_or_add_pPr()
    # Enlever jc existant
    for existing in pPr.findall(qn('w:jc')):
        pPr.remove(existing)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

def add_run(para, text, size_pt, bold=False, italic=False, color=None):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def fill_complet(data):
    doc = Document(io.BytesIO(CERT_COMPLET))
    civ = 'Madame' if data['civilite'] == 'F' else 'Monsieur'
    nom_complet = f"{data['prenom']} {data['nom']}"
    date_cours = datetime.strptime(data['date_cours'], '%Y-%m-%d').strftime('%d.%m.%Y')
    date_sig = datetime.today().strftime('%d.%m.%Y')

    # Para 8 = civilité (vide dans complet)
    para8 = doc.paragraphs[8]
    clear_para(para8)
    set_center(para8)
    add_run(para8, civ, 14, color=(0x5B, 0x5B, 0x5B))

    # Para 9 = nom (RecipientName, vide dans complet)
    para9 = doc.paragraphs[9]
    clear_para(para9)
    set_center(para9)
    add_run(para9, nom_complet, 20, bold=True, color=(0xC0, 0x39, 0x2B))

    # Para 14 = "Le :" + date
    para14 = doc.paragraphs[14]
    for run in para14.runs:
        if run.text.strip() == '':
            run.text = f' {date_cours}'
            break
    else:
        para14.add_run(f' {date_cours}')

    # Table row 0
    table = doc.tables[0]
    _fill_table(table, date_sig, data['formateur'])

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def fill_compact(data):
    doc = Document(io.BytesIO(CERT_COMPACT))
    civ = 'Madame' if data['civilite'] == 'F' else 'Monsieur'
    nom_complet = f"{data['prenom']} {data['nom']}"
    date_cours = datetime.strptime(data['date_cours'], '%Y-%m-%d').strftime('%d.%m.%Y')
    date_sig = datetime.today().strftime('%d.%m.%Y')

    # Para 5 = "Monsieur/Madame" → remplacer par civilité
    para5 = doc.paragraphs[5]
    clear_para(para5)
    set_center(para5)
    add_run(para5, civ, 14, color=(0x5B, 0x5B, 0x5B))

    # Para 7 = "Tartenpion marcel" → remplacer par nom
    para7 = doc.paragraphs[7]
    clear_para(para7)
    set_center(para7)
    add_run(para7, nom_complet, 20, bold=True, color=(0xC0, 0x39, 0x2B))

    # Para 12 = "Le : " → ajouter date
    para12 = doc.paragraphs[12]
    for run in para12.runs:
        if 'Le' in run.text:
            run.text = f'Le : {date_cours}'
            break

    # Table row 0
    table = doc.tables[0]
    _fill_table(table, date_sig, data['formateur'])

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def _fill_table(table, date_sig, formateur):
    # Row 0 cell 0 = date signature
    cell_date = table.rows[0].cells[0]
    for para in cell_date.paragraphs:
        clear_para(para)
    p = cell_date.paragraphs[0]
    set_center(p)
    add_run(p, date_sig, 11, italic=True, color=(0xC0, 0x39, 0x2B))

    # Row 0 cell 2 = formateur
    cell_sign = table.rows[0].cells[2]
    for para in cell_sign.paragraphs:
        clear_para(para)
    p2 = cell_sign.paragraphs[0]
    set_center(p2)
    r = add_run(p2, formateur, 13, italic=True, color=(0xC0, 0x39, 0x2B))
    r.font.name = 'Brush Script MT'

    # Vider row 1
    for ci in [0, 2]:
        for para in table.rows[1].cells[ci].paragraphs:
            clear_para(para)

def convert_to_pdf(docx_bytes):
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'cert.docx')
        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, docx_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            raise Exception(f'LibreOffice error: {result.stderr}')
        pdf_path = docx_path.replace('.docx', '.pdf')
        with open(pdf_path, 'rb') as f:
            return f.read()

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

@app.route('/generate-cert', methods=['POST'])
def generate_cert():
    try:
        data = request.json
        for field in ['prenom', 'nom', 'civilite', 'cours', 'date_cours', 'formateur']:
            if field not in data:
                return jsonify({'error': f'Champ manquant: {field}'}), 400

        if 'Complet' in data['cours']:
            docx_bytes = fill_complet(data)
        else:
            docx_bytes = fill_compact(data)

        pdf_bytes = convert_to_pdf(docx_bytes)
        filename = f"Certificat_{data['prenom']}_{data['nom']}.pdf"
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)

# ==================== EMAIL ====================

BREVO_API_KEY = os.environ.get('BREVO_API_KEY', '')

def send_email_formateur(formateur_email, formateur_nom, cours_data):
    """Envoyer un email de notification au formateur"""
    if not formateur_email:
        return False
    
    date_f = datetime.strptime(cours_data['date_cours'], '%Y-%m-%d').strftime('%d.%m.%Y')
    
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background: #c0392b; padding: 20px; text-align: center;">
            <h1 style="color: white; margin: 0; font-size: 24px;">SWISS ViTa Form</h1>
            <p style="color: rgba(255,255,255,0.85); margin: 5px 0 0 0;">Nouveau cours assigné</p>
        </div>
        <div style="padding: 30px; background: #f9f9f9;">
            <p>Bonjour {formateur_nom},</p>
            <p>Un nouveau cours vous a été assigné. Voici les informations :</p>
            <div style="background: white; border-radius: 8px; padding: 20px; margin: 20px 0; border-left: 4px solid #c0392b;">
                <table style="width: 100%; border-collapse: collapse;">
                    <tr><td style="padding: 8px 0; color: #888; width: 140px;">Type de cours</td><td style="padding: 8px 0; font-weight: bold;">{cours_data['type_cours']}</td></tr>
                    <tr><td style="padding: 8px 0; color: #888;">Date</td><td style="padding: 8px 0; font-weight: bold;">{date_f}</td></tr>
                    <tr><td style="padding: 8px 0; color: #888;">Horaire</td><td style="padding: 8px 0;">{cours_data['heure_debut']} – {cours_data['heure_fin']}</td></tr>
                    <tr><td style="padding: 8px 0; color: #888;">Lieu</td><td style="padding: 8px 0;">{cours_data['lieu']}</td></tr>
                    <tr><td style="padding: 8px 0; color: #888;">Participants prévus</td><td style="padding: 8px 0;">{cours_data.get('nb_participants_prevus', '—')}</td></tr>
                </table>
            </div>
            {f'<div style="background: #e6f1fb; border-radius: 8px; padding: 16px; margin: 16px 0;"><strong>Notes :</strong><br><pre style="font-family: Arial; white-space: pre-wrap; margin: 8px 0 0 0;">{cours_data["notes"]}</pre></div>' if cours_data.get('notes') else ''}
            <p>Veuillez confirmer votre présence en vous connectant au portail :</p>
            <div style="text-align: center; margin: 24px 0;">
                <a href="https://portail-swissvitaform.netlify.app" style="background: #c0392b; color: white; padding: 12px 28px; border-radius: 8px; text-decoration: none; font-weight: bold;">Accéder au portail</a>
            </div>
        </div>
        <div style="background: #f0f0f0; padding: 16px; text-align: center; font-size: 12px; color: #888;">
            Swiss ViTa Form — Av. Kiener 29, 1400 Yverdon-les-Bains — 078 892 02 63
        </div>
    </div>
    """
    
    payload = {
        "sender": {"name": "Swiss ViTa Form", "email": "info@swissvf.ch"},
        "to": [{"email": formateur_email, "name": formateur_nom}],
        "subject": f"Cours assigné : {cours_data['type_cours']} — {date_f}",
        "htmlContent": html_content
    }
    
    response = requests.post(
        'https://api.brevo.com/v3/smtp/email',
        headers={'api-key': BREVO_API_KEY, 'Content-Type': 'application/json'},
        json=payload
    )
    print(f'[BREVO formateur] status={response.status_code} body={response.text}')
    return response.status_code == 201, response.text

@app.route('/send-notification', methods=['POST'])
def send_notification():
    try:
        data = request.json
        formateur_email = data.get('formateur_email', '')
        formateur_nom = data.get('formateur_nom', '')
        cours = data.get('cours', {})
        
        if not formateur_email:
            return jsonify({'error': 'Email formateur manquant'}), 400
        
        success, detail = send_email_formateur(formateur_email, formateur_nom, cours)
        if success:
            return jsonify({'status': 'sent'})
        else:
            return jsonify({'error': 'Echec envoi email', 'detail': detail}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def send_email_client_welcome(client_email, client_nom, login, mot_de_passe):
    """Envoyer un email de bienvenue au nouveau client avec ses identifiants"""
    if not client_email:
        return False

    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background: #c0392b; padding: 20px; text-align: center;">
            <h1 style="color: white; margin: 0; font-size: 24px;">SWISS ViTa Form</h1>
            <p style="color: rgba(255,255,255,0.85); margin: 5px 0 0 0;">Bienvenue</p>
        </div>
        <div style="padding: 30px; background: #f9f9f9;">
            <p>Bonjour,</p>
            <p>Bienvenue chez Swiss ViTa Form ! Un espace client a été créé pour <strong>{client_nom}</strong> sur notre portail.</p>
            <p>Voici vos identifiants de connexion :</p>
            <div style="background: white; border-radius: 8px; padding: 20px; margin: 20px 0; border-left: 4px solid #c0392b;">
                <table style="width: 100%; border-collapse: collapse;">
                    <tr><td style="padding: 8px 0; color: #888; width: 140px;">Identifiant</td><td style="padding: 8px 0; font-weight: bold;">{login}</td></tr>
                    <tr><td style="padding: 8px 0; color: #888;">Mot de passe</td><td style="padding: 8px 0; font-weight: bold;">{mot_de_passe}</td></tr>
                </table>
            </div>
            <p>Vous pouvez vous connecter dès maintenant pour consulter vos cours à venir et télécharger les certificats de vos participants :</p>
            <div style="text-align: center; margin: 24px 0;">
                <a href="https://client-svf.netlify.app" style="background: #c0392b; color: white; padding: 12px 28px; border-radius: 8px; text-decoration: none; font-weight: bold;">Accéder au portail</a>
            </div>
        </div>
        <div style="background: #f0f0f0; padding: 16px; text-align: center; font-size: 12px; color: #888;">
            Swiss ViTa Form — Av. Kiener 29, 1400 Yverdon-les-Bains — 078 892 02 63
        </div>
    </div>
    """

    payload = {
        "sender": {"name": "Swiss ViTa Form", "email": "info@swissvf.ch"},
        "to": [{"email": client_email, "name": client_nom}],
        "subject": "Bienvenue chez Swiss ViTa Form — vos identifiants",
        "htmlContent": html_content
    }

    response = requests.post(
        'https://api.brevo.com/v3/smtp/email',
        headers={'api-key': BREVO_API_KEY, 'Content-Type': 'application/json'},
        json=payload
    )
    print(f'[BREVO client welcome] status={response.status_code} body={response.text}')
    return response.status_code == 201, response.text


@app.route('/send-welcome-client', methods=['POST'])
def send_welcome_client():
    try:
        data = request.json
        client_email = data.get('client_email', '')
        client_nom = data.get('client_nom', '')
        login = data.get('login', '')
        mot_de_passe = data.get('mot_de_passe', '')

        if not client_email:
            return jsonify({'error': 'Email client manquant'}), 400

        success, detail = send_email_client_welcome(client_email, client_nom, login, mot_de_passe)
        if success:
            return jsonify({'status': 'sent'})
        else:
            return jsonify({'error': 'Echec envoi email', 'detail': detail}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


WEBHOOK_SECRET = os.environ.get('WEBHOOK_SECRET', '')


def send_email_formateur_welcome(formateur_email, formateur_nom, login, mot_de_passe):
    """Envoyer un email de bienvenue au nouveau formateur avec ses identifiants"""
    if not formateur_email:
        return False, 'Email manquant'

    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background: #c0392b; padding: 20px; text-align: center;">
            <h1 style="color: white; margin: 0; font-size: 24px;">SWISS ViTa Form</h1>
            <p style="color: rgba(255,255,255,0.85); margin: 5px 0 0 0;">Bienvenue</p>
        </div>
        <div style="padding: 30px; background: #f9f9f9;">
            <p>Bonjour {formateur_nom},</p>
            <p>Bienvenue chez Swiss ViTa Form ! Un accès au portail formateur a été créé pour vous.</p>
            <p>Voici vos identifiants de connexion :</p>
            <div style="background: white; border-radius: 8px; padding: 20px; margin: 20px 0; border-left: 4px solid #c0392b;">
                <table style="width: 100%; border-collapse: collapse;">
                    <tr><td style="padding: 8px 0; color: #888; width: 140px;">Identifiant</td><td style="padding: 8px 0; font-weight: bold;">{login}</td></tr>
                    <tr><td style="padding: 8px 0; color: #888;">Mot de passe</td><td style="padding: 8px 0; font-weight: bold;">{mot_de_passe}</td></tr>
                </table>
            </div>
            <p>Vous pouvez vous connecter dès maintenant pour consulter vos cours assignés :</p>
            <div style="text-align: center; margin: 24px 0;">
                <a href="https://client-svf.netlify.app" style="background: #c0392b; color: white; padding: 12px 28px; border-radius: 8px; text-decoration: none; font-weight: bold;">Accéder au portail</a>
            </div>
        </div>
        <div style="background: #f0f0f0; padding: 16px; text-align: center; font-size: 12px; color: #888;">
            Swiss ViTa Form — Av. Kiener 29, 1400 Yverdon-les-Bains — 078 892 02 63
        </div>
    </div>
    """

    payload = {
        "sender": {"name": "Swiss ViTa Form", "email": "info@swissvf.ch"},
        "to": [{"email": formateur_email, "name": formateur_nom}],
        "subject": "Bienvenue chez Swiss ViTa Form — vos identifiants",
        "htmlContent": html_content
    }

    response = requests.post(
        'https://api.brevo.com/v3/smtp/email',
        headers={'api-key': BREVO_API_KEY, 'Content-Type': 'application/json'},
        json=payload
    )
    print(f'[BREVO formateur welcome] status={response.status_code} body={response.text}')
    return response.status_code == 201, response.text


@app.route('/webhook-new-formateur', methods=['POST'])
def webhook_new_formateur():
    try:
        # Vérification du secret partagé (header configuré côté Supabase)
        if WEBHOOK_SECRET:
            incoming_secret = request.headers.get('X-Webhook-Secret', '')
            if incoming_secret != WEBHOOK_SECRET:
                return jsonify({'error': 'Non autorisé'}), 401

        data = request.json or {}
        # Supabase envoie {"type":"INSERT","table":"formateurs","record":{...}}
        record = data.get('record', data)

        formateur_email = record.get('email', '')
        prenom = record.get('prenom', '')
        nom = record.get('nom', '')
        formateur_nom = f'{prenom} {nom}'.strip()
        login = record.get('login', '')
        mot_de_passe = record.get('mot_de_passe', '')

        if not formateur_email:
            return jsonify({'error': 'Email formateur manquant dans le webhook'}), 400

        success, detail = send_email_formateur_welcome(formateur_email, formateur_nom, login, mot_de_passe)
        if success:
            return jsonify({'status': 'sent'})
        else:
            return jsonify({'error': 'Echec envoi email', 'detail': detail}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500
